"""
Resume Parser Service
Extracts text and structure from various resume formats (PDF, DOCX, TXT)
"""

import os
import re
from typing import Dict, Any, List, Optional
import logging
from pathlib import Path

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import resumepy
    RESUMEPY_AVAILABLE = True
except ImportError:
    RESUMEPY_AVAILABLE = False

try:
    from services.markitdown_service import markitdown_service
    MARKITDOWN_AVAILABLE = markitdown_service.is_available()
except ImportError:
    MARKITDOWN_AVAILABLE = False

from utils_others.logger import logger

class ResumeParser:
    def __init__(self):
        self.supported_formats = ["pdf", "docx", "txt", "pptx", "xlsx", "png", "jpg", "jpeg", "gif", "bmp", "tiff", "html", "htm"]
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_pattern = re.compile(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
        self.use_markitdown = MARKITDOWN_AVAILABLE
        
    async def parse_resume(self, file_path: str, content_type: str) -> Dict[str, Any]:
        """
        Parse resume file and extract structured information
        """
        try:
            # Extract text based on file type
            text_content = await self._extract_text(file_path, content_type)
            
            # Parse structured information
            parsed_data = {
                "raw_text": text_content,
                "contact_info": self._extract_contact_info(text_content),
                "sections": self._identify_sections(text_content),
                "skills": self._extract_skills(text_content),
                "experience": self._extract_experience(text_content),
                "education": self._extract_education(text_content),
                "summary": self._extract_summary(text_content),
                "metadata": {
                    "file_type": content_type,
                    "text_length": len(text_content),
                    "word_count": len(text_content.split())
                }
            }
            
            logger.info(f"Successfully parsed resume: {len(text_content)} characters extracted")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Resume parsing failed: {str(e)}")
            raise Exception(f"Failed to parse resume: {str(e)}")
    
    async def _extract_text(self, file_path: str, content_type: str) -> str:
        """Extract text from different file formats"""
        try:
            # Try MarkItDown first if available (supports more formats and better extraction)
            if self.use_markitdown and markitdown_service.is_supported_format(file_path):
                logger.info(f"Using MarkItDown for {content_type}")
                markdown_text = markitdown_service.convert_to_markdown(file_path)
                if markdown_text:
                    return self._clean_text(markdown_text)
                logger.warning(f"MarkItDown failed, falling back to traditional parser")
            
            # Fallback to traditional parsers
            if content_type == "application/pdf":
                return await self._extract_from_pdf(file_path)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._extract_from_docx(file_path)
            elif content_type == "text/plain":
                return await self._extract_from_txt(file_path)
            elif content_type in ["application/vnd.openxmlformats-officedocument.presentationml.presentation",  # PPTX
                                 "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:  # XLSX
                # These are only supported by MarkItDown
                if self.use_markitdown:
                    markdown_text = markitdown_service.convert_to_markdown(file_path)
                    if markdown_text:
                        return self._clean_text(markdown_text)
                raise ValueError(f"Format {content_type} requires MarkItDown library")
            elif content_type in ["image/png", "image/jpeg", "image/jpg", "image/gif", "image/bmp", "image/tiff"]:
                # Image OCR via MarkItDown
                if self.use_markitdown:
                    markdown_text = markitdown_service.convert_to_markdown(file_path)
                    if markdown_text:
                        return self._clean_text(markdown_text)
                raise ValueError(f"Image format {content_type} requires MarkItDown library")
            elif content_type in ["text/html", "application/xhtml+xml"]:
                # HTML parsing
                if self.use_markitdown:
                    markdown_text = markitdown_service.convert_to_markdown(file_path)
                    if markdown_text:
                        return self._clean_text(markdown_text)
                raise ValueError(f"HTML format {content_type} requires MarkItDown library")
            else:
                raise ValueError(f"Unsupported content type: {content_type}")
        except Exception as e:
            logger.error(f"Text extraction failed for {content_type}: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF is not installed. Add 'PyMuPDF==1.24.9' to requirements.txt")
        
        try:
            text = ""
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                text += page_text + "\n"
            
            doc.close()
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {str(e)}")
            raise
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Add 'python-docx==1.1.2' to requirements.txt")
        
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            raise
    
    async def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return self._clean_text(text)
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return self._clean_text(text)
            except Exception as e:
                logger.error(f"TXT text extraction failed: {str(e)}")
                raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text while preserving structure"""
        # Remove excessive whitespace but preserve line breaks
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove special characters that might cause issues but keep important ones
        text = re.sub(r'[^\w\s\-.@,;:()/#&\'\n]', ' ', text)
        
        # Normalize multiple line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        
        return '\n'.join(lines).strip()
    
    def _extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information from text"""
        contact_info = {
            "emails": list(set(self.email_pattern.findall(text))),
            "phones": list(set(self.phone_pattern.findall(text))),
            "linkedin": self._extract_linkedin(text),
            "github": self._extract_github(text),
            "website": self._extract_website(text)
        }
        
        return contact_info
    
    def _extract_linkedin(self, text: str) -> List[str]:
        """Extract LinkedIn profiles"""
        pattern = r'linkedin\.com/in/[\w-]+'
        return list(set(re.findall(pattern, text, re.IGNORECASE)))
    
    def _extract_github(self, text: str) -> List[str]:
        """Extract GitHub profiles"""
        pattern = r'github\.com/[\w-]+'
        return list(set(re.findall(pattern, text, re.IGNORECASE)))
    
    def _extract_website(self, text: str) -> List[str]:
        """Extract personal websites"""
        pattern = r'https?://[^\s<>"]+|www\.[^\s<>"]+'
        matches = re.findall(pattern, text, re.IGNORECASE)
        # Filter out LinkedIn and GitHub
        filtered = [url for url in matches if not any(x in url.lower() for x in ['linkedin', 'github'])]
        return list(set(filtered))
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify different resume sections"""
        sections = {}
        
        # Common section headers - more comprehensive patterns
        section_patterns = {
            "summary": r'^(?:PROFESSIONAL\s+SUMMARY|SUMMARY|OBJECTIVE|PROFILE|PROFESSIONAL\s+PROFILE)$',
            "experience": r'^(?:EXPERIENCE|WORK\s+EXPERIENCE|PROFESSIONAL\s+EXPERIENCE|EMPLOYMENT|WORK\s+HISTORY)$',
            "education": r'^(?:EDUCATION|ACADEMIC\s+BACKGROUND|QUALIFICATIONS|EDUCATIONAL\s+BACKGROUND)$',
            "skills": r'^(?:SKILLS|TECHNICAL\s+SKILLS|COMPETENCIES|EXPERTISE|TECHNICAL\s+EXPERTISE)$',
            "projects": r'^(?:PROJECTS|PERSONAL\s+PROJECTS|ACADEMIC\s+PROJECTS|PROJECT\s+EXPERIENCE)$',
            "certifications": r'^(?:CERTIFICATIONS|CERTIFICATES|CREDENTIALS|PROFESSIONAL\s+CERTIFICATIONS)$',
            "languages": r'^(?:LANGUAGES|LANGUAGE\s+PROFICIENCY|LANGUAGE\s+SKILLS)$'
        }
        
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            section_found = False
            for section_name, pattern in section_patterns.items():
                if re.match(pattern, line, re.IGNORECASE):
                    # Save previous section
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    
                    # Start new section
                    current_section = section_name
                    current_content = []
                    section_found = True
                    break
            
            if not section_found and current_section:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text"""
        # Common technical skills
        tech_skills = [
            'python', 'java', 'javascript', 'react', 'node.js', 'sql', 'html', 'css',
            'aws', 'docker', 'kubernetes', 'git', 'linux', 'mongodb', 'postgresql',
            'machine learning', 'data analysis', 'tensorflow', 'pytorch', 'nlp'
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in tech_skills:
            if skill in text_lower:
                found_skills.append(skill)
        
        return list(set(found_skills))
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience information"""
        experiences = []
        
        # Look for patterns like "Company Name - Position (Year)"
        pattern = r'([A-Z][a-zA-Z\s&]+)\s*[-–]\s*([A-Z][a-zA-Z\s]+)\s*\((\d{4}[-–]?\d{0,4})\)'
        matches = re.findall(pattern, text)
        
        for match in matches:
            company, position, duration = match
            experiences.append({
                "company": company.strip(),
                "position": position.strip(),
                "duration": duration.strip()
            })
        
        return experiences
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information"""
        education = []
        
        # Look for degree patterns
        pattern = r'(Bachelor|Master|PhD|B\.S\.|M\.S\.|B\.A\.|M\.A\.)[^,\n]*,\s*([A-Z][a-zA-Z\s]+)[^,\n]*,\s*(\d{4})'
        matches = re.findall(pattern, text)
        
        for match in matches:
            degree, institution, year = match
            education.append({
                "degree": degree.strip(),
                "institution": institution.strip(),
                "year": year.strip()
            })
        
        return education
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary"""
        sections = self._identify_sections(text)
        return sections.get("summary", "").strip() if "summary" in sections else None
